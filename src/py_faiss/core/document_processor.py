import asyncio
import logging
from pathlib import Path
from typing import Union

import aiofiles
import chardet
import fitz
from docx import Document

from py_faiss.config import settings

import asyncio
import os
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import aiofiles
import tempfile
import hashlib
from datetime import datetime
import mimetypes

# 文档处理库
from docx import Document
import pandas as pd
from openpyxl import load_workbook
import chardet

from py_faiss.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """文档处理器 - 支持多种文档格式的文本提取和处理"""

    def __init__(self):
        self.supported_extensions = {
            '.docx', '.doc',  # Word 文档
            '.pdf',  # PDF 文档
            '.txt', '.md',  # 文本文档
            '.xlsx', '.xls',  # Excel 文档
            '.csv',  # CSV 文档
            '.json',  # JSON 文档
            '.xml',  # XML 文档
        }

        # 文本分割配置
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

        # 临时文件目录
        self.temp_dir = Path(settings.TEMP_PATH)
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    async def extract_text(self, file_path: Union[str, Path]):
        """
        从文档中提取文本

        Args:
            file_path: 文档文件路径

        Returns:
            提取的文本内容
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 检查文件大小
        file_size = file_path.stat().st_size
        if file_size > settings.MAX_FILE_SIZE:
            raise ValueError(f"文件过大: {file_size / 1024 / 1024:.1f}MB > {settings.MAX_FILE_SIZE / 1024 / 1024}MB")

        # 获取文件扩展名
        extension = file_path.suffix.lower()

        if extension not in self.supported_extensions:
            raise ValueError(f"不支持的文件格式: {extension}")

        try:
            # 根据文件类型选择提取方法
            if extension in ['.docx', '.doc']:
                text = await self._extract_from_docx(file_path)
            elif extension == '.pdf':
                text = await self._extract_from_pdf(file_path)
            elif extension in ['.txt', '.md']:
                text = await self._extract_from_text(file_path)
            elif extension in ['.xlsx', '.xls']:
                text = await self._extract_from_excel(file_path)
            elif extension == '.csv':
                text = await self._extract_from_csv(file_path)
            elif extension == '.json':
                text = await self._extract_from_json(file_path)
            elif extension == '.xml':
                text = await self._extract_from_xml(file_path)
            else:
                raise ValueError(f"暂不支持的文件格式: {extension}")

            logger.info(f"成功提取文本: {file_path.name}, 长度: {len(text)} 字符")
            return text

        except Exception as e:
            logger.error(f"文本提取失败 {file_path.name}: {e}")
            raise

    async def _extract_from_docx(self, file_path: Path) -> str:
        """从 DOCX 文件提取文本"""
        try:
            # 在线程池中执行 IO 密集型操作
            loop = asyncio.get_event_loop()

            def _read_docx():
                doc = Document(file_path)
                paragraphs = []

                # 提取段落文件
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        paragraphs.append(text)

                # 提取表格文本
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            paragraphs.append(' | '.join(row_text))

                return '\n'.join(paragraphs)

            text = await loop.run_in_executor(None, _read_docx)
            return text

        except Exception as e:
            raise Exception(f"DOCX 处理失败: {e}")

    async def _extract_from_pdf(self, file_path: Path) -> str:
        """从 PDF 文件提取文本"""
        try:
            loop = asyncio.get_event_loop()

            def _read_pdf():
                text_content = []

                with open(file_path, 'rb') as file:
                    doc = fitz.open(file)

                    return doc.name

            text = await loop.run_in_executor(None, _read_pdf)
            return text

        except Exception as e:
            raise Exception(f"PDF 处理失败: {e}")

    async def _extract_from_text(self, file_path: Path) -> str:
        """从文本文件提取内容"""
        try:
            # 检测文件编码
            async with aiofiles.open(file_path, 'rb') as f:
                raw_data = await f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'

            # 读取文本
            async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                text = await f.read()
                return text.strip()

        except Exception as e:
            # 尝试常见编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        text = await f.read()
                        return text.strip()
                except Exception as e:
                    logger.warning(f"尝试编码 {encoding} 失败: {e}")
                    continue

            raise Exception(f"文本文件读取失败: {e}")

    async def _extract_from_excel(self, file_path: Path) -> str:
        """从 Excel 文件提取文本"""
        try:
            loop = asyncio.get_event_loop()

            def _read_excel():
                content = []

                if file_path.suffix.lower() == '.xlsx':
                    # 使用 openpyxl 读取 .xlsx
                    wb = load_workbook(file_path, read_only=True)
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        content.append(f"[工作表: {sheet_name}]")

                        for row in sheet.iter_rows(values_only=True):
                            row_data = [str(cell) if cell is not None else '' for cell in row]
                            row_text = ' | '.join(filter(None, row_data))
                            if row_text.strip():
                                content.append(row_text)
                        content.append("")
                else:
                    # 使用 pandas 读取 .xls
                    excel_file = pd.ExcelFile(file_path)
                    for sheet_name in excel_file.sheet_names:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        content.append(f"[工作表: {sheet_name}]")

                        # 添加列标题
                        headers = ' | '.join(str(col) for col in df.columns)
                        content.append(headers)

                        # 添加数据行
                        for _, row in df.iterrows():
                            row_data = [str(val) if pd.notna(val) else '' for val in row]
                            row_text = ' | '.join(filter(None, row_data))
                            if row_text.strip():
                                content.append(row_text)
                        content.append("")

                return '\n'.join(content)

            text = await loop.run_in_executor(None, _read_excel)
            return text

        except Exception as e:
            raise Exception(f"Excel 处理失败: {e}")

    async def _extract_from_csv(self, file_path: Path) -> str:
        """从 CSV 文件提取文本"""
        try:
            loop = asyncio.get_event_loop()

            def _read_csv():
                # 尝试不同编码
                for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding)
                        break
                    except:
                        continue
                else:
                    raise Exception("无法识别 CSV 文件编码")

                content = []

                # 添加列标题
                headers = ' | '.join(str(col) for col in df.columns)
                content.append(headers)

                # 添加数据行
                for _, row in df.iterrows():
                    row_data = [str(val) if pd.notna(val) else '' for val in row]
                    row_text = ' | '.join(filter(None, row_data))
                    if row_text.strip():
                        content.append(row_text)

                return '\n'.join(content)

            text = await loop.run_in_executor(None, _read_csv)
            return text

        except Exception as e:
            raise Exception(f"CSV 处理失败: {e}")

    async def _extract_from_json(self, file_path: Path) -> str:
        """从 JSON 文件提取文本"""
        try:
            import json

            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                data = json.loads(content)

                # 递归提取 JSON 中的文本内容
                def extract_text_from_json(obj, path=""):
                    texts = []

                    if isinstance(obj, dict):
                        for key, value in obj.items():
                            current_path = f"{path}.{key}" if path else key
                            texts.extend(extract_text_from_json(value, current_path))
                    elif isinstance(obj, list):
                        for i, item in enumerate(obj):
                            current_path = f"{path}[{i}]" if path else f"[{i}]"
                            texts.extend(extract_text_from_json(item, current_path))
                    else:
                        # 基本类型，转换为字符串
                        text = str(obj).strip()
                        if text and len(text) > 1:  # 过滤单字符
                            texts.append(f"{path}: {text}")

                    return texts

                text_parts = extract_text_from_json(data)
                return '\n'.join(text_parts)

        except Exception as e:
            raise Exception(f"JSON 处理失败: {e}")

    async def _extract_from_xml(self, file_path: Path) -> str:
        """从 XML 文件提取文本"""
        try:
            import xml.etree.ElementTree as ET

            loop = asyncio.get_event_loop()

            def _read_xml():
                tree = ET.parse(file_path)
                root = tree.getroot()

                def extract_xml_text(element, level=0):
                    texts = []
                    indent = "  " * level

                    # 元素标签和属性
                    if element.attrib:
                        attrs = ", ".join(f"{k}={v}" for k, v in element.attrib.items())
                        texts.append(f"{indent}<{element.tag} {attrs}>")
                    else:
                        texts.append(f"{indent}<{element.tag}>")

                    # 元素文本内容
                    if element.text and element.text.strip():
                        texts.append(f"{indent}  {element.text.strip()}")

                    # 递归处理子元素
                    for child in element:
                        texts.extend(extract_xml_text(child, level + 1))

                    return texts

                text_parts = extract_xml_text(root)
                return '\n'.join(text_parts)

            text = await loop.run_in_executor(None, _read_xml)
            return text

        except Exception as e:
            raise Exception(f"XML 处理失败: {e}")